"""Unit tests for the document conversion pipeline."""

from __future__ import annotations

import sys
from pathlib import Path

import pytest

_BACKEND_ROOT = Path(__file__).resolve().parents[1]
if str(_BACKEND_ROOT) not in sys.path:
    sys.path.insert(0, str(_BACKEND_ROOT))

from app.ingestion import docx_to_markdown


# ── Docx conversion ─────────────────────────────────────────────────────────


class TestDocxToMarkdown:
    def test_converts_simple_docx(self):
        """Create a minimal docx in memory and verify conversion."""
        from docx import Document

        import io
        doc = Document()
        doc.add_heading("Test Heading", level=1)
        doc.add_paragraph("Body text here.")
        doc.add_heading("Sub Heading", level=2)
        doc.add_paragraph("More content.")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        md = docx_to_markdown(buf.read())
        assert "# Test Heading" in md
        assert "## Sub Heading" in md
        assert "Body text here." in md
        assert "More content." in md

    def test_handles_missing_paragraph_style(self, monkeypatch):
        """Paragraphs with missing style metadata should still convert."""
        from types import SimpleNamespace

        fake_doc = SimpleNamespace(
            paragraphs=[
                SimpleNamespace(text="Styled heading", style=SimpleNamespace(name="Heading 1")),
                SimpleNamespace(text="Plain paragraph", style=None),
                SimpleNamespace(text="No style name", style=SimpleNamespace()),
            ]
        )

        monkeypatch.setattr("app.ingestion.DocxDocument", lambda _: fake_doc)

        md = docx_to_markdown(b"fake-docx")
        assert "# Styled heading" in md
        assert "Plain paragraph" in md
        assert "No style name" in md
